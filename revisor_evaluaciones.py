"""
SISTEMA REVISOR DE EVALUACIONES
Basado en Matriz Técnico-Pedagógica de Accesibilidad Evaluativa
Versión: 1.0 — 3 Categorías / 17 Criterios
Uso: python revisor_evaluaciones.py <archivo_evaluacion.docx> [tiempo_en_minutos]
"""

import sys
import re
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# DICCIONARIOS DE DETECCIÓN
# ============================================================

CONECTORES_SUBORDINANTES = [
    'aunque', 'cuando', 'mientras', 'porque', 'debido a que', 'ya que',
    'puesto que', 'dado que', 'si bien', 'a pesar de que', 'con tal de que',
    'para que', 'sin que', 'después de que', 'antes de que', 'siempre que',
    'a menos que', 'de modo que', 'de manera que'
]

CONECTORES_COMPLEJOS = [
    'no obstante', 'por consiguiente', 'en virtud de', 'sin embargo',
    'en consecuencia', 'por ende', 'de ahí que', 'con todo', 'ahora bien',
    'así pues', 'en efecto', 'a tal efecto', 'en virtud de lo anterior'
]

VERBOS_AMBIGUOS = [
    'reflexiona', 'comenta', 'desarrolla', 'analiza', 'considera',
    'explora', 'revisa', 'trabaja', 'trata', 'aborda', 'observa',
    'piensa', 'imagina', 'discute', 'estudia', 'examina'
]

VERBOS_ALTA_DEMANDA = [
    'argumenta', 'justifica', 'evalúa', 'critica', 'fundamenta',
    'sintetiza', 'diseña', 'formula', 'propone', 'defiende', 'contrasta',
    'argumente', 'justifique', 'evalúe', 'fundamente', 'sintetice'
]

VERBOS_IMPERATIVOS = [
    'analiza', 'responde', 'justifica', 'compara', 'describe', 'explica',
    'define', 'menciona', 'calcula', 'identifica', 'clasifica', 'lista',
    'enumera', 'demuestra', 'aplica', 'resuelve', 'evalúa', 'sintetiza',
    'argumenta', 'fundamenta', 'critica', 'propone', 'desarrolla', 'comenta',
    'reflexiona', 'señala', 'indica', 'escribe', 'completa', 'une', 'marca',
    'elige', 'selecciona', 'ordena', 'relaciona', 'nombra', 'redacta',
    'analice', 'responda', 'justifique', 'compare', 'describa', 'explique',
    'defina', 'mencione', 'calcule', 'identifique', 'clasifique', 'liste',
    'enumere', 'demuestre', 'aplique', 'resuelva', 'evalúe', 'sintetice'
]

TERMINOS_PUNITIVOS = [
    'obligatoriamente', 'sin excepción', 'bajo ninguna circunstancia',
    'estrictamente', 'es obligatorio', 'no puede dejar', 'es exigido'
]

PALABRAS_SANCION = [
    'se descontará', 'será anulado', 'perderá puntos', 'descuento',
    'penalización', 'anulará', 'invalidará', 'se restará'
]

PALABRAS_VELOCIDAD = [
    'rápidamente', 'de inmediato', 'sin demorarse', 'a la brevedad',
    'velozmente', 'lo antes posible', 'urgente'
]

# ============================================================
# EXTRACCIÓN DE TEXTO
# ============================================================

def extraer_texto(filepath):
    doc = Document(filepath)
    parrafos = [p.text for p in doc.paragraphs]
    texto_completo = '\n'.join(parrafos)
    return parrafos, texto_completo

def obtener_oraciones(texto):
    oraciones = re.split(r'[.!?;]', texto)
    return [o.strip() for o in oraciones if o.strip()]

def obtener_preguntas(parrafos):
    preguntas = []
    for p in parrafos:
        if '?' in p or re.match(r'^\s*\d+[\.\)]', p.strip()):
            preguntas.append(p)
    return preguntas

# ============================================================
# FUNCIONES DE ANÁLISIS — 17 CRITERIOS
# ============================================================

def analizar_1_1(texto, parrafos):
    oraciones = obtener_oraciones(texto)
    oraciones_largas = [o for o in oraciones if len(o.split()) > 30]
    multi_verbo = []
    for o in oraciones:
        verbs = [v for v in VERBOS_IMPERATIVOS if re.search(r'\b' + v + r'\b', o.lower())]
        if len(verbs) >= 3:
            multi_verbo.append(o)

    resultado = {'id': '1.1', 'nombre': 'Extensión excesiva de enunciados', 'hallazgos': [], 'estado': 'OK'}
    if oraciones_largas:
        resultado['hallazgos'].append(f'{len(oraciones_largas)} enunciado(s) con más de 30 palabras continuas sin puntuación fuerte.')
        resultado['estado'] = 'ALERTA'
    if multi_verbo:
        resultado['hallazgos'].append(f'{len(multi_verbo)} enunciado(s) con 3 o más verbos de acción acumulados (Ej: "describe, compara y argumenta").')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Dividir enunciados largos en segmentos breves. Secuenciar las acciones en pasos numerados.'
    return resultado

def analizar_1_2(texto, parrafos):
    oraciones = obtener_oraciones(texto)
    voz_pasiva = len(re.findall(r'\b(será|serán|fue|fueron|es|son)\s+\w+[aeiou]d[oa]\b', texto.lower()))
    conectores_hallados = [c for c in CONECTORES_COMPLEJOS if c in texto.lower()]
    muchas_comas = [o for o in oraciones if o.count(',') > 3]

    resultado = {'id': '1.2', 'nombre': 'Complejidad sintáctica', 'hallazgos': [], 'estado': 'OK'}
    if voz_pasiva >= 2:
        resultado['hallazgos'].append(f'{voz_pasiva} estructura(s) en voz pasiva detectada(s) (ej: "será desarrollado"). Usar voz activa.')
        resultado['estado'] = 'ALERTA'
    if len(conectores_hallados) >= 3:
        resultado['hallazgos'].append(f'Conectores complejos: {", ".join(conectores_hallados[:5])}. Sustituir por conectores simples.')
        resultado['estado'] = 'ALERTA'
    elif conectores_hallados:
        resultado['hallazgos'].append(f'Conectores complejos presentes: {", ".join(conectores_hallados)}.')
        resultado['estado'] = 'REVISAR'
    if muchas_comas:
        resultado['hallazgos'].append(f'{len(muchas_comas)} oración(es) con más de 3 comas. Fragmentar el enunciado.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Reescribir en oraciones simples. Reemplazar voz pasiva por voz activa. Usar conectores simples (y, pero, entonces).'
    return resultado

def analizar_1_3(texto, parrafos):
    ambiguos = [v for v in VERBOS_AMBIGUOS if re.search(r'\b' + v + r'\b', texto.lower())]
    tiene_cantidad = bool(re.search(r'(menciona|explica|describe|nombra|señala)\s+\d+', texto.lower()))
    sin_limite = [p for p in parrafos if any(v in p.lower() for v in ['analiza el', 'comenta el', 'desarrolla']) and not re.search(r'\d+', p)]

    resultado = {'id': '1.3', 'nombre': 'Ambigüedad en instrucción', 'hallazgos': [], 'estado': 'OK'}
    if ambiguos:
        resultado['hallazgos'].append(f'Verbos ambiguos detectados: {", ".join(ambiguos)}. Reemplazar por verbos observables (define, menciona, calcula).')
        resultado['estado'] = 'ALERTA'
    if sin_limite:
        resultado['hallazgos'].append(f'{len(sin_limite)} consigna(s) abierta(s) sin delimitación de extensión o cantidad esperada.')
        resultado['estado'] = 'ALERTA'
    if not tiene_cantidad and len(parrafos) > 5:
        resultado['hallazgos'].append('No se detectan delimitadores concretos ("menciona 2", "explica 1"). Considerar agregar.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Usar verbos observables y precisar la cantidad o extensión de respuesta esperada en cada ítem.'
    return resultado

def analizar_1_4(texto, parrafos):
    bloques_largos = 0
    consecutivos = 0
    for p in parrafos:
        if p.strip():
            consecutivos += 1
            if consecutivos >= 6:
                bloques_largos += 1
        else:
            consecutivos = 0
    sin_numeracion = not bool(re.search(r'^\s*\d+[\.\)]|^\s*[a-z]\)', texto, re.MULTILINE))
    multi_accion_sin_lista = [p for p in parrafos if len([v for v in VERBOS_IMPERATIVOS if v in p.lower()]) >= 3 and not re.match(r'^\s*\d+|^\s*[a-z]\)', p.strip())]

    resultado = {'id': '1.4', 'nombre': 'Densidad de información visual', 'hallazgos': [], 'estado': 'OK'}
    if bloques_largos > 0:
        resultado['hallazgos'].append(f'Se detectaron bloques de 6 o más líneas consecutivas sin salto visual.')
        resultado['estado'] = 'ALERTA'
    if sin_numeracion and len(parrafos) > 8:
        resultado['hallazgos'].append('No se detecta numeración ni viñetas en las instrucciones.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    if multi_accion_sin_lista:
        resultado['hallazgos'].append(f'{len(multi_accion_sin_lista)} instrucción(es) con múltiples acciones sin estructura de lista.')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Dividir en bloques breves. Incorporar numeración, viñetas y espacios visuales entre preguntas.'
    return resultado

def analizar_1_5(texto, parrafos):
    tecnicismos = ['hipótesis', 'inferencia', 'paradigma', 'epistemológico', 'fenomenológico',
                   'heurístico', 'axioma', 'silogismo', 'tautología', 'diacrónico', 'ontológico']
    hallados = [t for t in tecnicismos if t in texto.lower()]
    adverbios_mente = [w for w in re.findall(r'\b\w+mente\b', texto.lower()) if len(w) > 12]
    nominalizaciones = [n for n in re.findall(r'\b\w+(?:ción|miento|ncia|idad)\b', texto.lower()) if len(n) > 10]
    polisilabos = [w for w in re.findall(r'\b[a-záéíóúñ]{13,}\b', texto.lower())]

    resultado = {'id': '1.5', 'nombre': 'Vocabulario de alta complejidad', 'hallazgos': [], 'estado': 'OK'}
    if len(hallados) >= 4:
        resultado['hallazgos'].append(f'Tecnicismos sin apoyo contextual: {", ".join(hallados[:6])}. Incorporar glosario.')
        resultado['estado'] = 'ALERTA'
    if len(adverbios_mente) >= 3:
        resultado['hallazgos'].append(f'Adverbios complejos en "-mente": {", ".join(set(adverbios_mente[:4]))}.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    if len(nominalizaciones) >= 6:
        resultado['hallazgos'].append(f'Alta presencia de nominalizaciones abstractas ({len(nominalizaciones)} detectadas).')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    if len(polisilabos) >= 5:
        resultado['hallazgos'].append(f'Palabras de gran longitud detectadas: {", ".join(set(polisilabos[:4]))}...')
    resultado['sugerencia'] = 'Incorporar glosario breve o contextualizar términos técnicos. Simplificar vocabulario cuando sea posible.'
    return resultado

def analizar_1_6(texto, parrafos):
    pronombres_ambiguos = re.findall(r'\b(esto|ello|aquello)\s+(demuestra|indica|muestra|significa|implica)\b', texto.lower())
    dobles_negaciones = re.findall(r'\bno\b.{0,20}\b(incorrecta|incorrecto|falsa|ninguno|jamás|tampoco)\b', texto.lower())
    preguntas_negativas = [p for p in parrafos if re.search(r'¿.{0,60}(NO|no es|no corresponde|no sería).{0,30}\?', p)]
    cadenas_causales = [p for p in parrafos if len(re.findall(r'\b(si|porque|debido a|ya que|puesto que)\b', p.lower())) >= 3]

    resultado = {'id': '1.6', 'nombre': 'Inferencia excesiva requerida', 'hallazgos': [], 'estado': 'OK'}
    if len(pronombres_ambiguos) >= 2:
        resultado['hallazgos'].append(f'Pronombres ambiguos sin antecedente cercano ({len(pronombres_ambiguos)} casos). Reemplazar por referente explícito.')
        resultado['estado'] = 'ALERTA'
    if dobles_negaciones:
        resultado['hallazgos'].append(f'{len(dobles_negaciones)} doble(s) negación detectada(s). Reformular en afirmativo.')
        resultado['estado'] = 'ALERTA'
    if preguntas_negativas:
        resultado['hallazgos'].append(f'{len(preguntas_negativas)} pregunta(s) formuladas con negación explícita (¿Cuál NO...?).')
        resultado['estado'] = 'ALERTA'
    if cadenas_causales:
        resultado['hallazgos'].append(f'{len(cadenas_causales)} pregunta(s) con múltiples conectores causales acumulados.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Reemplazar pronombres por referentes explícitos. Reformular negaciones. Simplificar estructura causal.'
    return resultado

def analizar_2_1(texto, parrafos, tiempo_minutos=None):
    preguntas = obtener_preguntas(parrafos)
    n_preguntas = len(preguntas)
    desarrollo_extenso = [p for p in parrafos if any(v in p.lower() for v in ['explique', 'justifique', 'fundamente', 'argumente', 'desarrolle'])]
    multi_tarea = [p for p in parrafos if len([v for v in VERBOS_IMPERATIVOS if v in p.lower()]) >= 4]

    resultado = {'id': '2.1', 'nombre': 'Cantidad excesiva de ítems', 'hallazgos': [], 'estado': 'OK'}
    if tiempo_minutos and n_preguntas > 0:
        if n_preguntas / tiempo_minutos > 0.5:
            resultado['hallazgos'].append(f'Alta relación ítems/tiempo: {n_preguntas} preguntas detectadas para {tiempo_minutos} minutos (supera 1 ítem cada 2 min).')
            resultado['estado'] = 'ALERTA'
    elif n_preguntas >= 20:
        resultado['hallazgos'].append(f'Se detectaron {n_preguntas} preguntas/ítems. Verificar si el tiempo asignado es suficiente.')
        resultado['estado'] = 'REVISAR'
    if len(desarrollo_extenso) > 3:
        resultado['hallazgos'].append(f'{len(desarrollo_extenso)} preguntas de desarrollo extenso detectadas (recomendado: máximo 3).')
        resultado['estado'] = 'ALERTA'
    if multi_tarea:
        resultado['hallazgos'].append(f'{len(multi_tarea)} ítem(s) con 4 o más microtareas acumuladas en una sola pregunta.')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Reducir número de ítems o aumentar tiempo. Priorizar profundidad sobre cantidad. Dividir tareas complejas.'
    return resultado

def analizar_2_2(texto, parrafos):
    textos_largos = re.findall(r'(?:Lee|Observa|Analiza|Considera el siguiente).{200,}', texto, re.DOTALL)
    sobrecarga_proceso = [p for p in parrafos if sum(1 for v in ['leer', 'calcular', 'justificar', 'analizar', 'comparar', 'redactar', 'graficar'] if v in p.lower()) >= 3]

    resultado = {'id': '2.2', 'nombre': 'Tiempo insuficiente por tarea', 'hallazgos': [], 'estado': 'OK'}
    if textos_largos:
        resultado['hallazgos'].append(f'{len(textos_largos)} texto(s) de lectura extenso(s) detectado(s). Verificar proporcionalidad con preguntas asociadas.')
        resultado['estado'] = 'REVISAR'
    if sobrecarga_proceso:
        resultado['hallazgos'].append(f'{len(sobrecarga_proceso)} pregunta(s) exigen más de 3 procesos cognitivos simultáneos (leer + calcular + justificar).')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Fragmentar lecturas largas. Aumentar tiempo o simplificar tareas. Secuenciar procesos cognitivos en etapas.'
    return resultado

def analizar_2_3(texto, parrafos):
    items_complejos = [(p, [v for v in VERBOS_IMPERATIVOS if v in p.lower()]) for p in parrafos if len([v for v in VERBOS_IMPERATIVOS if v in p.lower()]) >= 3]

    resultado = {'id': '2.3', 'nombre': 'Multiplicidad de demandas simultáneas', 'hallazgos': [], 'estado': 'OK'}
    if items_complejos:
        resultado['hallazgos'].append(f'{len(items_complejos)} ítem(s) con 3 o más instrucciones operativas simultáneas en una sola pregunta.')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Separar habilidades en etapas. Dividir instrucciones en pasos numerados independientes.'
    return resultado

def analizar_2_4(texto, parrafos):
    parrafos_con_contenido = [p for p in parrafos if p.strip()]
    parrafos_vacios = [p for p in parrafos if not p.strip()]
    total = len(parrafos_con_contenido) + len(parrafos_vacios)

    resultado = {'id': '2.4', 'nombre': 'Formato visual saturado', 'hallazgos': [], 'estado': 'OK'}
    if total > 0:
        densidad = len(parrafos_con_contenido) / total
        if densidad > 0.85:
            resultado['hallazgos'].append(f'Alta densidad visual: muy pocas líneas en blanco (densidad {round(densidad*100)}%). Incorporar espacios.')
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Aumentar márgenes, espacios entre preguntas y líneas en blanco. Simplificar distribución por página.'
    return resultado

def analizar_2_5(texto, parrafos):
    tipos_formato = set()
    for p in parrafos:
        pl = p.lower().strip()
        if re.search(r'\bseleccione\b|\bmarque\b|\bencierre\b|[a-d]\)', pl):
            tipos_formato.add('selección múltiple')
        if re.search(r'explique|justifique|desarrolle|argumente|fundamente', pl):
            tipos_formato.add('desarrollo')
        if re.search(r'complete|llene|_____|rellene', pl):
            tipos_formato.add('completar')
        if re.search(r'relacione|una|conecte con flechas', pl):
            tipos_formato.add('unir/relacionar')
        if re.search(r'verdadero|falso|V o F|V/F', pl):
            tipos_formato.add('V/F')

    resultado = {'id': '2.5', 'nombre': 'Cambios frecuentes de formato', 'hallazgos': [], 'estado': 'OK'}
    if len(tipos_formato) >= 4:
        resultado['hallazgos'].append(f'Se detectaron {len(tipos_formato)} tipos distintos de formato de respuesta: {", ".join(tipos_formato)}.')
        resultado['estado'] = 'ALERTA'
    elif len(tipos_formato) >= 3:
        resultado['hallazgos'].append(f'{len(tipos_formato)} formatos de respuesta detectados: {", ".join(tipos_formato)}. Considerar agrupar.')
        resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Agrupar ítems del mismo formato en secciones. Incorporar títulos de sección. Mantener estructura visual estable.'
    return resultado

def analizar_3_1(texto, parrafos):
    punitivos = [t for t in TERMINOS_PUNITIVOS if t in texto.lower()]
    sanciones = [t for t in PALABRAS_SANCION if t in texto.lower()]
    mayusculas_exceso = re.findall(r'[A-ZÁÉÍÓÚÑ]{5,}', texto)
    velocidad = [t for t in PALABRAS_VELOCIDAD if t in texto.lower()]
    neg_directives = re.findall(r'\bNO\s+[a-záéíóúñ]+', texto)

    resultado = {'id': '3.1', 'nombre': 'Lenguaje amenazante o punitivo', 'hallazgos': [], 'estado': 'OK'}
    if punitivos:
        resultado['hallazgos'].append(f'Términos de presión u obligación: {", ".join(punitivos)}. Sustituir por lenguaje neutral.')
        resultado['estado'] = 'ALERTA'
    if sanciones:
        resultado['hallazgos'].append(f'Advertencias con sanción detectadas: {", ".join(sanciones)}. Eliminar o reformular.')
        resultado['estado'] = 'ALERTA'
    if len(mayusculas_exceso) >= 4:
        resultado['hallazgos'].append(f'Uso excesivo de mayúsculas completas ({len(mayusculas_exceso)} instancias). Reducir énfasis agresivo.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    if velocidad:
        resultado['hallazgos'].append(f'Lenguaje de presión de velocidad: {", ".join(velocidad)}.')
        resultado['estado'] = 'ALERTA'
    if len(neg_directives) > 5:
        resultado['hallazgos'].append(f'{len(neg_directives)} instrucciones prohibitivas ("NO + verbo"). Reformular en positivo.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Sustituir lenguaje punitivo por formulaciones neutras y positivas. Eliminar advertencias innecesarias.'
    return resultado

def analizar_3_2(parrafos):
    primeros = [p for p in parrafos if p.strip()][:12]
    complejas_al_inicio = [p for p in primeros if any(re.search(r'\b' + v + r'\b', p.lower()) for v in VERBOS_ALTA_DEMANDA)]
    desarrollo_al_inicio = [p for p in primeros if re.search(r'explique detalladamente|fundamente ampliamente|desarrolle en profundidad', p.lower())]
    palabras_inicio = sum(len(p.split()) for p in primeros)

    resultado = {'id': '3.2', 'nombre': 'Alta concentración de dificultad inicial', 'hallazgos': [], 'estado': 'OK'}
    if len(complejas_al_inicio) >= 2:
        resultado['hallazgos'].append(f'{len(complejas_al_inicio)} pregunta(s) de alta demanda cognitiva en los primeros ítems. Iniciar con preguntas más accesibles.')
        resultado['estado'] = 'ALERTA'
    if desarrollo_al_inicio:
        resultado['hallazgos'].append('Preguntas de desarrollo extenso detectadas en los primeros ítems. Incorporar preguntas breves iniciales.')
        resultado['estado'] = 'ALERTA'
    if palabras_inicio > 300:
        resultado['hallazgos'].append(f'Alta densidad textual al inicio ({palabras_inicio} palabras en los primeros bloques). Distribuir carga progresivamente.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Iniciar con preguntas de reconocimiento o respuesta breve. Distribuir complejidad en forma ascendente.'
    return resultado

def analizar_3_3(texto, parrafos):
    tiene_encabezados = bool(re.search(r'(parte|sección|bloque|ítem)\s+[IVX\d]', texto, re.IGNORECASE))
    sin_ejemplo = [p for p in parrafos if any(v in p.lower() for v in VERBOS_AMBIGUOS) and 'ejemplo' not in p.lower() and 'ej.' not in p.lower()]

    resultado = {'id': '3.3', 'nombre': 'Instrucciones poco predecibles', 'hallazgos': [], 'estado': 'OK'}
    if not tiene_encabezados and len(parrafos) > 15:
        resultado['hallazgos'].append('No se detectan encabezados de sección (Parte I, Sección A). El formato puede ser impredecible.')
        resultado['estado'] = 'REVISAR'
    if len(sin_ejemplo) >= 3:
        resultado['hallazgos'].append(f'{len(sin_ejemplo)} instrucción(es) ambigua(s) sin ejemplo orientador.')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Incorporar encabezados entre secciones. Agregar ejemplos en formatos complejos. Mantener patrón de respuesta estable.'
    return resultado

def analizar_3_4(texto, parrafos):
    preguntas_negativas = [p for p in parrafos if re.search(r'¿.{0,80}(NO|no es|no corresponde|no sería).{0,30}\?', p)]
    dobles_neg = re.findall(r'\bno\b.{0,20}\b(incorrecta|incorrecto|ninguno|jamás)\b', texto.lower())
    evitar_errores = [t for t in ['evite equivocarse', 'no cometa errores', 'tenga cuidado de no', 'evite errores'] if t in texto.lower()]

    resultado = {'id': '3.4', 'nombre': 'Penalización excesiva del error', 'hallazgos': [], 'estado': 'OK'}
    if preguntas_negativas:
        resultado['hallazgos'].append(f'{len(preguntas_negativas)} pregunta(s) formulada(s) en negativo (¿Cuál NO...?). Reformular en afirmativo.')
        resultado['estado'] = 'ALERTA'
    if dobles_neg:
        resultado['hallazgos'].append(f'{len(dobles_neg)} doble(s) negación detectada(s). Simplificar redacción.')
        resultado['estado'] = 'ALERTA'
    if evitar_errores:
        resultado['hallazgos'].append(f'Indicaciones centradas en evitar errores: "{", ".join(evitar_errores)}". Reformular hacia logro esperado.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Reformular preguntas negativas en afirmativo. Enfocar instrucciones en el logro, no en el error.'
    return resultado

def analizar_3_5(texto, parrafos):
    refs_tiempo = re.findall(r'\d+\s*minutos?|\d+\s*min\.?', texto.lower())
    velocidad = [t for t in PALABRAS_VELOCIDAD if t in texto.lower()]
    items_compactos = [p for p in parrafos if re.match(r'^\s*\d+[\.\)]', p.strip()) and len(p.split()) < 15]

    resultado = {'id': '3.5', 'nombre': 'Sobrecarga temporal explícita', 'hallazgos': [], 'estado': 'OK'}
    if len(refs_tiempo) > 4:
        resultado['hallazgos'].append(f'Referencias temporales repetidas: {len(refs_tiempo)} menciones de tiempo. Unificar en tiempo global.')
        resultado['estado'] = 'ALERTA'
    if velocidad:
        resultado['hallazgos'].append(f'Lenguaje de urgencia/rapidez detectado: {", ".join(velocidad)}.')
        resultado['estado'] = 'ALERTA'
    if len(items_compactos) > 15:
        resultado['hallazgos'].append(f'{len(items_compactos)} preguntas breves y consecutivas sin pausa visual. Incorporar espacios.')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    resultado['sugerencia'] = 'Reducir presión temporal explícita. Unificar tiempos globales. Priorizar calidad sobre velocidad.'
    return resultado

def analizar_3_6(texto, parrafos):
    primeros = [p for p in parrafos if p.strip()][:6]
    tiene_intro = any(re.search(r'(instruccion|orientacion|a continuacion|estimado|esta evaluacion|objetivo)', p.lower()) for p in primeros)
    tiene_secciones = bool(re.search(r'(parte|sección|bloque)\s+[IVX\d]', texto[:600], re.IGNORECASE))
    comienza_con_numeros = bool(re.match(r'^\s*\d+[\.\)]', '\n'.join(primeros[:3])))

    resultado = {'id': '3.6', 'nombre': 'Escasa orientación inicial', 'hallazgos': [], 'estado': 'OK'}
    if not tiene_intro:
        resultado['hallazgos'].append('No se detecta bloque de orientación o instrucción general al inicio.')
        resultado['estado'] = 'ALERTA'
    if not tiene_secciones and len(parrafos) > 20:
        resultado['hallazgos'].append('No se detecta estructura visible por secciones (sin títulos o numeración de partes).')
        if resultado['estado'] == 'OK':
            resultado['estado'] = 'REVISAR'
    if comienza_con_numeros:
        resultado['hallazgos'].append('La evaluación comienza directamente con preguntas numeradas, sin orientación previa.')
        resultado['estado'] = 'ALERTA'
    resultado['sugerencia'] = 'Incorporar bloque introductorio con instrucciones generales. Organizar por secciones con títulos claros.'
    return resultado

# ============================================================
# GENERACIÓN DEL INFORME WORD
# ============================================================

ESTADO_COLOR = {
    'OK':      RGBColor(0x37, 0x56, 0x23),
    'REVISAR': RGBColor(0x7F, 0x60, 0x00),
    'ALERTA':  RGBColor(0xC0, 0x39, 0x0B),
}
ESTADO_BG = {
    'OK':      '92D050',
    'REVISAR': 'FFD966',
    'ALERTA':  'FF6B35',
}
ESTADO_ETIQUETA = {
    'OK':      'SIN PROBLEMAS',
    'REVISAR': 'REVISAR',
    'ALERTA':  'ALERTA',
}

def set_bg_celda(celda, color_hex):
    tc = celda._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def p_titulo(doc, texto, pt=18, color=RGBColor(0x1F, 0x4E, 0x79), antes=8, despues=4):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold = True
    r.font.size = Pt(pt)
    r.font.color.rgb = color
    r.font.name = 'Arial'
    par.paragraph_format.space_before = Pt(antes)
    par.paragraph_format.space_after = Pt(despues)
    return par

def p_cuerpo(doc, texto, pt=11, negrita=False, italica=False, color=None, indent=False):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.font.size = Pt(pt)
    r.font.name = 'Arial'
    r.bold = negrita
    r.italic = italica
    if color:
        r.font.color.rgb = color
    if indent:
        par.paragraph_format.left_indent = Inches(0.3)
    par.paragraph_format.space_after = Pt(3)
    return par

def linea_separadora(doc, color='2E75B6'):
    par = doc.add_paragraph()
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    par.paragraph_format.space_after = Pt(4)
    return par

def generar_informe(resultados, archivo_fuente, archivo_salida, tiempo_minutos=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    # === PORTADA ===
    tit = doc.add_paragraph()
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tit.add_run('INFORME DE REVISIÓN TÉCNICO-PEDAGÓGICA')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79); r.font.name = 'Arial'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('Revisión Automática con Matriz de Accesibilidad Evaluativa')
    r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6); r2.font.name = 'Arial'

    doc.add_paragraph()

    # Tabla metadata
    t_meta = doc.add_table(rows=3, cols=2)
    t_meta.style = 'Table Grid'
    meta = [
        ('Archivo evaluado:', os.path.basename(archivo_fuente)),
        ('Fecha de revisión:', datetime.datetime.now().strftime('%d/%m/%Y  %H:%M hrs.')),
        ('Tiempo declarado:', f'{tiempo_minutos} minutos' if tiempo_minutos else 'No especificado'),
    ]
    for i, (lbl, val) in enumerate(meta):
        t_meta.rows[i].cells[0].text = lbl
        t_meta.rows[i].cells[1].text = val
        t_meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_bg_celda(t_meta.rows[i].cells[0], 'DEEAF1')

    doc.add_paragraph()

    # === RESUMEN ===
    alertas = [r for r in resultados if r['estado'] == 'ALERTA']
    revisar = [r for r in resultados if r['estado'] == 'REVISAR']
    ok_list = [r for r in resultados if r['estado'] == 'OK']

    p_titulo(doc, 'RESUMEN EJECUTIVO', 14)
    t_sum = doc.add_table(rows=4, cols=2)
    t_sum.style = 'Table Grid'
    datos_sum = [
        ('Criterios con ALERTA', str(len(alertas)), 'FF6B35'),
        ('Criterios a REVISAR', str(len(revisar)), 'FFD966'),
        ('Criterios sin problemas', str(len(ok_list)), '92D050'),
        ('Total criterios evaluados', str(len(resultados)), 'DEEAF1'),
    ]
    for i, (lbl, val, col) in enumerate(datos_sum):
        t_sum.rows[i].cells[0].text = lbl
        t_sum.rows[i].cells[1].text = val
        t_sum.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        t_sum.rows[i].cells[1].paragraphs[0].runs[0].bold = True
        t_sum.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_bg_celda(t_sum.rows[i].cells[0], col)
        set_bg_celda(t_sum.rows[i].cells[1], col)

    doc.add_paragraph()

    # === TABLA GENERAL DE CRITERIOS ===
    p_titulo(doc, 'ESTADO GENERAL — 17 CRITERIOS', 14)
    t_ov = doc.add_table(rows=len(resultados) + 1, cols=3)
    t_ov.style = 'Table Grid'
    for i, h in enumerate(['ID', 'Criterio', 'Estado']):
        c = t_ov.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_bg_celda(c, '1F4E79')
    for i, r in enumerate(resultados):
        row = t_ov.rows[i + 1]
        row.cells[0].text = r['id']
        row.cells[1].text = r['nombre']
        row.cells[2].text = ESTADO_ETIQUETA[r['estado']]
        row.cells[2].paragraphs[0].runs[0].bold = True
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_bg_celda(row.cells[2], ESTADO_BG[r['estado']])

    doc.add_page_break()

    # === DETALLE POR CATEGORÍA ===
    categorias = [
        ('CATEGORÍA 1: COMPRENSIÓN LECTORA', [r for r in resultados if r['id'].startswith('1.')], '1F4E79'),
        ('CATEGORÍA 2: VELOCIDAD DE PROCESAMIENTO', [r for r in resultados if r['id'].startswith('2.')], '2E75B6'),
        ('CATEGORÍA 3: ANSIEDAD EVALUATIVA', [r for r in resultados if r['id'].startswith('3.')], 'C0390B'),
    ]

    for titulo_cat, resultados_cat, color_hex in categorias:
        color_rgb = RGBColor(int(color_hex[:2],16), int(color_hex[2:4],16), int(color_hex[4:],16))
        par_cat = doc.add_paragraph()
        r_cat = par_cat.add_run(titulo_cat)
        r_cat.bold = True; r_cat.font.size = Pt(15); r_cat.font.color.rgb = color_rgb; r_cat.font.name = 'Arial'
        par_cat.paragraph_format.space_before = Pt(10); par_cat.paragraph_format.space_after = Pt(4)
        pPr = par_cat._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        btm = OxmlElement('w:bottom')
        btm.set(qn('w:val'), 'single'); btm.set(qn('w:sz'), '8'); btm.set(qn('w:space'), '1'); btm.set(qn('w:color'), color_hex)
        pBdr.append(btm); pPr.append(pBdr)

        for res in resultados_cat:
            # Header criterio
            pc = doc.add_paragraph()
            rc = pc.add_run(f'Criterio {res["id"]} — {res["nombre"]}')
            rc.bold = True; rc.font.size = Pt(12); rc.font.name = 'Arial'
            pc.paragraph_format.space_before = Pt(8); pc.paragraph_format.space_after = Pt(2)

            # Estado
            pe = doc.add_paragraph()
            re_s = pe.add_run(f'  ESTADO: {ESTADO_ETIQUETA[res["estado"]]}  ')
            re_s.bold = True; re_s.font.size = Pt(11); re_s.font.color.rgb = ESTADO_COLOR[res['estado']]; re_s.font.name = 'Arial'
            pe.paragraph_format.space_after = Pt(4)

            # Hallazgos
            if res['hallazgos']:
                ph = doc.add_paragraph()
                rh = ph.add_run('Hallazgos detectados:')
                rh.bold = True; rh.font.size = Pt(11); rh.font.color.rgb = RGBColor(0x40, 0x40, 0x40); rh.font.name = 'Arial'
                for hallazgo in res['hallazgos']:
                    pb = doc.add_paragraph(style='List Bullet')
                    rb = pb.add_run(hallazgo)
                    rb.font.size = Pt(11); rb.font.name = 'Arial'
                    pb.paragraph_format.left_indent = Inches(0.3)
                    pb.paragraph_format.space_after = Pt(2)
            else:
                pok = doc.add_paragraph()
                rok = pok.add_run('No se detectaron problemas en este criterio.')
                rok.italic = True; rok.font.size = Pt(11); rok.font.color.rgb = RGBColor(0x37, 0x56, 0x23); rok.font.name = 'Arial'

            # Sugerencia
            ps = doc.add_paragraph()
            rs1 = ps.add_run('Sugerencia: ')
            rs1.bold = True; rs1.font.size = Pt(11); rs1.font.color.rgb = color_rgb; rs1.font.name = 'Arial'
            rs2 = ps.add_run(res['sugerencia'])
            rs2.italic = True; rs2.font.size = Pt(11); rs2.font.color.rgb = color_rgb; rs2.font.name = 'Arial'
            ps.paragraph_format.space_after = Pt(10)

        doc.add_paragraph()

    # === NOTA METODOLÓGICA ===
    doc.add_page_break()
    p_titulo(doc, 'NOTA METODOLÓGICA', 13, RGBColor(0x59, 0x59, 0x59))
    notas = [
        'Informe generado automáticamente aplicando la Matriz Técnico-Pedagógica de Accesibilidad Evaluativa.',
        'Criterios organizados en 3 categorías: Comprensión Lectora (1.1–1.6), Velocidad de Procesamiento (2.1–2.5) y Ansiedad Evaluativa (3.1–3.6).',
        'Los estados ALERTA y REVISAR son indicadores orientadores. El docente debe contextualizar los hallazgos.',
        'Escala: VERDE = sin problemas  /  AMARILLO = requiere revisión  /  NARANJA = problema detectado.',
        'Se recomienda atender primero los criterios ALERTA, luego los REVISAR.',
        'La detección es automática basada en patrones textuales. Algunos hallazgos pueden requerir confirmación manual.',
    ]
    for nota in notas:
        pn = doc.add_paragraph(style='List Bullet')
        rn = pn.add_run(nota)
        rn.font.size = Pt(10); rn.font.name = 'Arial'; rn.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.save(archivo_salida)
    return archivo_salida

# ============================================================
# FUNCIÓN PRINCIPAL
# ============================================================

def revisar_evaluacion(ruta_docx, ruta_salida=None, tiempo_minutos=None):
    if not os.path.exists(ruta_docx):
        print(f'ERROR: No se encuentra el archivo: {ruta_docx}')
        return None

    if ruta_salida is None:
        base = os.path.splitext(ruta_docx)[0]
        ruta_salida = base + '_INFORME_REVISION.docx'

    print(f'\n{"="*55}')
    print(f'  SISTEMA REVISOR DE EVALUACIONES')
    print(f'{"="*55}')
    print(f'  Archivo: {os.path.basename(ruta_docx)}')
    if tiempo_minutos:
        print(f'  Tiempo declarado: {tiempo_minutos} minutos')
    print(f'{"="*55}\n')

    parrafos, texto = extraer_texto(ruta_docx)
    print(f'  Texto extraído: {len(parrafos)} párrafos, {len(texto.split())} palabras\n')
    print('  Analizando 17 criterios...')

    resultados = [
        analizar_1_1(texto, parrafos),
        analizar_1_2(texto, parrafos),
        analizar_1_3(texto, parrafos),
        analizar_1_4(texto, parrafos),
        analizar_1_5(texto, parrafos),
        analizar_1_6(texto, parrafos),
        analizar_2_1(texto, parrafos, tiempo_minutos),
        analizar_2_2(texto, parrafos),
        analizar_2_3(texto, parrafos),
        analizar_2_4(texto, parrafos),
        analizar_2_5(texto, parrafos),
        analizar_3_1(texto, parrafos),
        analizar_3_2(parrafos),
        analizar_3_3(texto, parrafos),
        analizar_3_4(texto, parrafos),
        analizar_3_5(texto, parrafos),
        analizar_3_6(texto, parrafos),
    ]

    alertas = sum(1 for r in resultados if r['estado'] == 'ALERTA')
    revisar = sum(1 for r in resultados if r['estado'] == 'REVISAR')
    ok_n = sum(1 for r in resultados if r['estado'] == 'OK')

    print(f'\n  RESULTADO:')
    print(f'  ALERTAS  : {alertas} criterios')
    print(f'  REVISAR  : {revisar} criterios')
    print(f'  SIN PROB.: {ok_n} criterios')
    print(f'\n  Generando informe Word...')

    ruta_final = generar_informe(resultados, ruta_docx, ruta_salida, tiempo_minutos)
    print(f'\n  INFORME GUARDADO EN:')
    print(f'  {ruta_final}')
    print(f'{"="*55}\n')
    return ruta_final

# ============================================================
if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('\nUSO:')
        print('  python revisor_evaluaciones.py <archivo.docx> [tiempo_en_minutos]')
        print('\nEJEMPLOS:')
        print('  python revisor_evaluaciones.py "C:/Desktop/prueba.docx"')
        print('  python revisor_evaluaciones.py "C:/Desktop/prueba.docx" 60')
        sys.exit(1)

    ruta = sys.argv[1]
    tiempo = int(sys.argv[2]) if len(sys.argv) >= 3 else None
    revisar_evaluacion(ruta, tiempo_minutos=tiempo)
